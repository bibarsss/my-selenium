import re
from docx import Document
import os
from docx.shared import Pt
from docx.oxml.ns import qn

def sanitize(name: str) -> str:
    return re.sub(r'[<>:"/\\|?*]', '', name.strip())

def set_times_new_roman(run):
    run.font.name = "Times New Roman"
    run.font.italic = False         # 🔥 important
    # run.font.bold = False           # optional
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

def withoutGroup(data: dict, worker_id):
    output_dir = "generated_files"
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, sanitize(data['generated_file_name']))

    doc = Document(data['template_file_name'])

    for paragraph in doc.paragraphs:
        for key, val in data['replace'].items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, val)
                # override formatting of the new run
                if paragraph.runs:
                    set_times_new_roman(paragraph.runs[0])

    doc.save(output_path)

# def withoutGroupp(data: dict, worker_id):
#     output_dir = "generated_files"
#     os.makedirs(output_dir, exist_ok=True)
#     output_path = os.path.join(output_dir, data['generated_file_name'])

#     doc = Document(data['template_file_name'])

    # for paragraph in doc.paragraphs:
    #     for key, val in data['replace'].items():
    #         if key in paragraph.text:
    #             paragraph.text = paragraph.text.replace(key, val)

    # for paragraph in doc.paragraphs:
    #     text = paragraph.text

    #     for r in paragraph.runs:
    #         r.clear()

    #     new_run = paragraph.add_run(text)
    #     set_times_new_roman(new_run)

    # for paragraph in doc.paragraphs:
    #     text = paragraph.text

    #     for key, val in data['replace'].items():
    #         text = text.replace(key, val)

    #     for r in paragraph.runs:
    #         r.clear()

    #     new_run = paragraph.add_run(text)
    #     set_times_new_roman(new_run)

    # for table in doc.tables:
    #     for row in table.rows:
    #         for cell in row.cells:
    #             for paragraph in cell.paragraphs:
    #                 text = paragraph.text

    #                 for key, val in data['replace'].items():
    #                     text = text.replace(key, val)

    #                 for r in paragraph.runs:
    #                     r.clear()

    #                 new_run = paragraph.add_run(text)
    #                 set_times_new_roman(new_run)

    # doc.save(output_path)

def withGroup(data: list, worker_id):
    output_dir = "generated_files"
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, sanitize(data['generated_file_name']))

    new_replace = data['replace'][0].copy()
    for key_template, template in data['replace_template'].items():
        new_replace[key_template] = ''

        # count = 0
        for replace in data['replace']:
            # count += 1
            temp = template
            for k, v in replace.items():
                temp = temp.replace(k, v)
            new_replace[key_template] += temp
            # new_replace[key_template] += f"{count}. {temp}"

    doc = Document(data['template_file_name'])
    for paragraph in doc.paragraphs:
        for key, val in new_replace.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, val)
                # override formatting of the new run
                if paragraph.runs:
                    set_times_new_roman(paragraph.runs[0])

    # for paragraph in doc.paragraphs:
    #     text = paragraph.text

    #     for key, val in new_replace.items():
    #         text = text.replace(key, val)

    #     for r in paragraph.runs:
    #         r.clear()

    #     new_run = paragraph.add_run(text)
    #     set_times_new_roman(new_run)

    # for table in doc.tables:
    #     for row in table.rows:
    #         for cell in row.cells:
    #             for paragraph in cell.paragraphs:
    #                 text = paragraph.text

    #                 for key, val in data['replace'].items():
    #                     text = text.replace(key, val)

    #                 for r in paragraph.runs:
    #                     r.clear()

    #                 new_run = paragraph.add_run(text)
    #                 set_times_new_roman(new_run)

    doc.save(output_path)
